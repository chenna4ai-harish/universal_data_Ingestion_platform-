"""
file_parser.py
--------------
Handles all file ingestion: ZIP unpacking (nested), file-type detection,
encoding detection, and parsing to DataFrames.

Returns a list of ParsedFile dataclasses for downstream processing.
"""

from __future__ import annotations

import io
import os
import re
import tempfile
import unicodedata
import uuid
import zipfile
from dataclasses import dataclass, field
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

import pandas as pd


# ---------------------------------------------------------------------------
# Result types
# ---------------------------------------------------------------------------

@dataclass
class ParsedFile:
    """One successfully parsed source file."""
    source_filename: str
    dataframe: pd.DataFrame
    file_format: str          # csv / tsv / xlsx / json / xml / txt / html / docx / pdf_text / pdf_scanned_ocr
    encoding_used: str | None
    archive_lineage_id: str | None = None   # set if came from ZIP
    archive_meta: dict = field(default_factory=dict)


@dataclass
class FailedFile:
    """A file that could not be parsed or was blocked."""
    source_filename: str
    exception_type: str       # from global taxonomy
    reason: str
    archive_lineage_id: str | None = None
    archive_meta: dict = field(default_factory=dict)


@dataclass
class ArchiveLineageRow:
    """One row for the ARCHIVE_LINEAGE system table."""
    archive_lineage_id: str
    job_id: str
    source_filename: str
    parent_archive: str | None
    root_archive: str | None
    archive_entry_name: str
    extracted_path: str
    nested_level: int
    file_size_bytes: int
    extraction_status: str
    insert_timestamp: str


# ---------------------------------------------------------------------------
# Normalisation helpers
# ---------------------------------------------------------------------------

def _normalise_text(text: str, cfg: dict) -> str:
    """Apply unicode normalisation and control-char stripping per config."""
    norm = cfg.get("text_extraction", {}).get("unicode_normalization", "NFKC")
    text = unicodedata.normalize(norm, text)
    if cfg.get("text_extraction", {}).get("strip_control_characters", True):
        text = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", text)
    if cfg.get("text_extraction", {}).get("collapse_repeated_whitespace", True):
        text = re.sub(r"[ \t]+", " ", text)
    return text


def _detect_encoding(raw: bytes, cfg: dict) -> tuple[str, str]:
    """Try encoding fallback chain. Returns (decoded_str, encoding_used)."""
    order = cfg.get("text_extraction", {}).get(
        "encoding_detection_order",
        ["utf-8-sig", "utf-16", "utf-8", "cp1252", "latin-1"]
    )
    for enc in order:
        try:
            return raw.decode(enc), enc
        except (UnicodeDecodeError, LookupError):
            continue
    # last-resort latin-1 (never fails)
    return raw.decode("latin-1", errors="replace"), "latin-1"


# ---------------------------------------------------------------------------
# Individual format parsers
# ---------------------------------------------------------------------------

def _parse_csv_like(raw: bytes, cfg: dict, filename: str) -> tuple[pd.DataFrame, str]:
    """Parse CSV / TSV / pipe / semicolon delimited files."""
    text, enc = _detect_encoding(raw, cfg)
    text = _normalise_text(text, cfg)

    # Detect delimiter
    first_line = text.split("\n")[0] if "\n" in text else text[:500]
    counts = {
        ",": first_line.count(","),
        "\t": first_line.count("\t"),
        "|": first_line.count("|"),
        ";": first_line.count(";"),
    }
    delimiter = max(counts, key=counts.get)
    if counts[delimiter] == 0:
        delimiter = ","

    df = pd.read_csv(io.StringIO(text), sep=delimiter, dtype=str,
                     keep_default_na=False, on_bad_lines="skip")
    df.columns = [str(c).strip() for c in df.columns]
    ext = Path(filename).suffix.lower()
    fmt = {"tsv": "tsv", "txt": "txt"}.get(ext.lstrip("."), "csv")
    if delimiter == "\t":
        fmt = "tsv"
    elif delimiter == "|":
        fmt = "pipe_delimited"
    elif delimiter == ";":
        fmt = "semicolon_delimited"
    return df, enc


def _parse_excel(raw: bytes, cfg: dict, filename: str) -> tuple[pd.DataFrame, str]:
    ext = Path(filename).suffix.lower()
    engine = "openpyxl" if ext == ".xlsx" else "xlrd"
    df = pd.read_excel(io.BytesIO(raw), dtype=str, keep_default_na=False, engine=engine)
    df.columns = [str(c).strip() for c in df.columns]
    return df, "binary"


def _parse_json(raw: bytes, cfg: dict, filename: str) -> tuple[pd.DataFrame, str]:
    text, enc = _detect_encoding(raw, cfg)
    text = _normalise_text(text, cfg)
    df = pd.read_json(io.StringIO(text), dtype=str)
    if df.empty:
        raise ValueError("JSON parsed to empty DataFrame")
    df.columns = [str(c).strip() for c in df.columns]
    return df, enc


def _parse_xml(raw: bytes, cfg: dict, filename: str) -> tuple[pd.DataFrame, str]:
    text, enc = _detect_encoding(raw, cfg)
    text = _normalise_text(text, cfg)
    # pandas read_xml requires lxml or etree
    df = pd.read_xml(io.StringIO(text), dtype=str)
    df.columns = [str(c).strip() for c in df.columns]
    return df, enc


def _parse_html(raw: bytes, cfg: dict, filename: str) -> tuple[pd.DataFrame, str]:
    text, enc = _detect_encoding(raw, cfg)
    text = _normalise_text(text, cfg)
    tables = pd.read_html(io.StringIO(text), dtype_backend="numpy_nullable")
    if not tables:
        raise ValueError("No tables found in HTML")
    df = tables[0].astype(str)
    df.columns = [str(c).strip() for c in df.columns]
    return df, enc


def _parse_docx(raw: bytes, cfg: dict, filename: str) -> tuple[pd.DataFrame, str]:
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx not installed; run: pip install python-docx")

    doc = Document(io.BytesIO(raw))
    rows = []
    headers = None
    for table in doc.tables:
        for i, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            if i == 0 and headers is None:
                headers = cells
            else:
                rows.append(cells)
    if headers and rows:
        df = pd.DataFrame(rows, columns=headers, dtype=str)
        df.columns = [str(c).strip() for c in df.columns]
        return df, "utf-8"

    # Fall back: extract all paragraph text as single-column
    lines = [p.text for p in doc.paragraphs if p.text.strip()]
    df = pd.DataFrame({"text": lines}, dtype=str)
    return df, "utf-8"


def _parse_pdf(raw: bytes, cfg: dict, filename: str) -> tuple[pd.DataFrame, str]:
    """Extract text from PDF (text layer). Falls back to OCR if needed."""
    min_chars = cfg.get("text_extraction", {}).get("minimum_extracted_characters", 20)
    max_garbled = cfg.get("text_extraction", {}).get("max_garbled_character_ratio", 0.15)

    extracted_text = ""
    method = "pdf_text"

    try:
        import pdfplumber
        with pdfplumber.open(io.BytesIO(raw)) as pdf:
            for page in pdf.pages:
                t = page.extract_text()
                if t:
                    extracted_text += t + "\n"
    except ImportError:
        raise ImportError("pdfplumber not installed; run: pip install pdfplumber")

    # Garbled ratio check
    if extracted_text:
        non_print = sum(1 for c in extracted_text if unicodedata.category(c) in ("Cc", "Cs"))
        garbled = non_print / max(len(extracted_text), 1)
        if garbled > max_garbled:
            extracted_text = ""

    # OCR fallback
    if len(extracted_text.strip()) < min_chars:
        if cfg.get("text_extraction", {}).get("ocr_enabled_for_scanned_documents", True):
            try:
                import pytesseract
                from PIL import Image
                import pdfplumber
                pages_text = []
                with pdfplumber.open(io.BytesIO(raw)) as pdf:
                    for page in pdf.pages:
                        img = page.to_image(resolution=150).original
                        pages_text.append(pytesseract.image_to_string(img))
                extracted_text = "\n".join(pages_text)
                method = "pdf_scanned_ocr"
            except ImportError:
                raise ImportError(
                    "pytesseract / Pillow not installed for OCR; "
                    "run: pip install pytesseract Pillow"
                )

    if len(extracted_text.strip()) < min_chars:
        action = cfg.get("text_extraction", {}).get("empty_text_action", "route_to_exception")
        if action == "route_to_exception":
            raise ValueError(f"PDF extracted fewer than {min_chars} characters")

    extracted_text = _normalise_text(extracted_text, cfg)

    # Try to parse as CSV-like (common for PDF AR exports)
    lines = [l.strip() for l in extracted_text.splitlines() if l.strip()]
    if lines:
        try:
            df = pd.read_csv(io.StringIO("\n".join(lines)), dtype=str,
                             keep_default_na=False, on_bad_lines="skip")
            df.columns = [str(c).strip() for c in df.columns]
            if len(df.columns) > 1:
                return df, method
        except Exception:
            pass

    df = pd.DataFrame({"text": lines}, dtype=str)
    return df, method


# ---------------------------------------------------------------------------
# Format dispatcher
# ---------------------------------------------------------------------------

_TEXT_EXTENSIONS = {".csv", ".tsv", ".txt", ".tab"}
_EXCEL_EXTENSIONS = {".xlsx", ".xls"}
_BLOCKED_EXTENSIONS: set[str] = set()   # populated from config


def _get_file_format(filename: str) -> str | None:
    ext = Path(filename).suffix.lower()
    if ext in _TEXT_EXTENSIONS:
        return "text"
    if ext in _EXCEL_EXTENSIONS:
        return "excel"
    if ext == ".json":
        return "json"
    if ext == ".xml":
        return "xml"
    if ext == ".html" or ext == ".htm":
        return "html"
    if ext == ".docx":
        return "docx"
    if ext == ".pdf":
        return "pdf"
    return None


def _parse_single_file(
    raw: bytes,
    filename: str,
    cfg: dict,
) -> tuple[pd.DataFrame, str]:
    """Dispatch to appropriate parser. Returns (DataFrame, format_label)."""
    fmt = _get_file_format(filename)
    if fmt is None:
        raise ValueError(f"Unsupported file type: {Path(filename).suffix}")

    if fmt == "text":
        df, enc = _parse_csv_like(raw, cfg, filename)
        ext = Path(filename).suffix.lower()
        label = {"txt": "txt", "tsv": "tsv", "tab": "tsv"}.get(ext.lstrip("."), "csv")
        return df, label
    if fmt == "excel":
        df, enc = _parse_excel(raw, cfg, filename)
        return df, "xlsx" if filename.lower().endswith(".xlsx") else "xls"
    if fmt == "json":
        df, enc = _parse_json(raw, cfg, filename)
        return df, "json"
    if fmt == "xml":
        df, enc = _parse_xml(raw, cfg, filename)
        return df, "xml"
    if fmt == "html":
        df, enc = _parse_html(raw, cfg, filename)
        return df, "html"
    if fmt == "docx":
        df, enc = _parse_docx(raw, cfg, filename)
        return df, "docx"
    if fmt == "pdf":
        df, enc = _parse_pdf(raw, cfg, filename)
        return df, enc  # enc is 'pdf_text' or 'pdf_scanned_ocr' here
    raise ValueError(f"No parser for format: {fmt}")


# ---------------------------------------------------------------------------
# ZIP handling
# ---------------------------------------------------------------------------

def _is_blocked_extension(name: str, cfg: dict) -> bool:
    blocked = [e.lower() for e in cfg.get("ingestion", {}).get("blocked_extensions", [])]
    return Path(name).suffix.lower() in blocked


def _is_supported_extension(name: str, cfg: dict) -> bool:
    return _get_file_format(name) is not None


def _safe_extract_path(entry_name: str, extract_dir: str) -> str | None:
    """Return safe absolute path or None if zip-slip detected."""
    extract_dir = os.path.realpath(extract_dir)
    target = os.path.realpath(os.path.join(extract_dir, entry_name))
    if not target.startswith(extract_dir + os.sep) and target != extract_dir:
        return None  # zip-slip
    return target


def _unpack_zip(
    zip_bytes: bytes,
    zip_filename: str,
    cfg: dict,
    job_id: str,
    parsed_files: list[ParsedFile],
    failed_files: list[FailedFile],
    archive_lineage_rows: list[ArchiveLineageRow],
    parent_archive: str | None,
    root_archive: str | None,
    current_depth: int,
    file_count_tracker: list[int],   # mutable counter [total_files_so_far]
    extract_dir: str,
) -> None:
    """Recursively unpack a ZIP, respecting all safety controls."""
    ing = cfg.get("ingestion", {})
    max_depth = ing.get("max_archive_nesting_depth", 5)
    max_files = ing.get("max_files_per_archive", 5000)
    max_total_mb = ing.get("max_uncompressed_size_mb", 2048)
    max_single_mb = ing.get("max_single_extracted_file_mb", 500)
    block_encrypted = ing.get("block_encrypted_archives", True)
    detect_slip = ing.get("detect_zip_slip", True)
    quarantine = ing.get("quarantine_failed_entries", True)
    ts = datetime.now(timezone.utc).isoformat()

    if current_depth > max_depth:
        failed_files.append(FailedFile(
            source_filename=zip_filename,
            exception_type="ARCHIVE_ERROR",
            reason=f"Exceeded max_archive_nesting_depth ({max_depth})",
        ))
        return

    try:
        zf = zipfile.ZipFile(io.BytesIO(zip_bytes))
    except zipfile.BadZipFile as e:
        failed_files.append(FailedFile(
            source_filename=zip_filename,
            exception_type="ARCHIVE_ERROR",
            reason=f"Bad ZIP file: {e}",
        ))
        return

    # Check for encryption
    if block_encrypted:
        for info in zf.infolist():
            if info.flag_bits & 0x1:  # encrypted flag
                failed_files.append(FailedFile(
                    source_filename=zip_filename,
                    exception_type="ARCHIVE_ERROR",
                    reason="Encrypted archive blocked by policy",
                ))
                return

    for info in zf.infolist():
        entry_name = info.filename

        # Skip directory entries
        if entry_name.endswith("/"):
            continue

        file_count_tracker[0] += 1
        if file_count_tracker[0] > max_files:
            failed_files.append(FailedFile(
                source_filename=entry_name,
                exception_type="ARCHIVE_ERROR",
                reason=f"Exceeded max_files_per_archive ({max_files})",
            ))
            continue

        # Zip-slip check
        safe_path = _safe_extract_path(entry_name, extract_dir) if detect_slip else os.path.join(extract_dir, entry_name)
        alid = str(uuid.uuid4())

        if safe_path is None:
            archive_lineage_rows.append(ArchiveLineageRow(
                archive_lineage_id=alid, job_id=job_id,
                source_filename=os.path.basename(entry_name),
                parent_archive=zip_filename, root_archive=root_archive or zip_filename,
                archive_entry_name=entry_name, extracted_path="BLOCKED",
                nested_level=current_depth, file_size_bytes=info.file_size,
                extraction_status="BLOCKED_ZIP_SLIP", insert_timestamp=ts,
            ))
            failed_files.append(FailedFile(
                source_filename=entry_name, exception_type="ARCHIVE_ERROR",
                reason="ZIP-slip path traversal blocked",
                archive_lineage_id=alid,
            ))
            continue

        # Blocked extension check
        if _is_blocked_extension(entry_name, cfg):
            archive_lineage_rows.append(ArchiveLineageRow(
                archive_lineage_id=alid, job_id=job_id,
                source_filename=os.path.basename(entry_name),
                parent_archive=zip_filename, root_archive=root_archive or zip_filename,
                archive_entry_name=entry_name, extracted_path=safe_path,
                nested_level=current_depth, file_size_bytes=info.file_size,
                extraction_status="BLOCKED_EXTENSION", insert_timestamp=ts,
            ))
            failed_files.append(FailedFile(
                source_filename=entry_name, exception_type="ARCHIVE_ERROR",
                reason=f"Blocked file extension: {Path(entry_name).suffix}",
                archive_lineage_id=alid,
            ))
            continue

        # Size check
        if info.file_size > max_single_mb * 1024 * 1024:
            archive_lineage_rows.append(ArchiveLineageRow(
                archive_lineage_id=alid, job_id=job_id,
                source_filename=os.path.basename(entry_name),
                parent_archive=zip_filename, root_archive=root_archive or zip_filename,
                archive_entry_name=entry_name, extracted_path=safe_path,
                nested_level=current_depth, file_size_bytes=info.file_size,
                extraction_status="BLOCKED_SIZE", insert_timestamp=ts,
            ))
            failed_files.append(FailedFile(
                source_filename=entry_name, exception_type="ARCHIVE_ERROR",
                reason=f"Entry size {info.file_size} exceeds limit {max_single_mb}MB",
                archive_lineage_id=alid,
            ))
            continue

        # Extract bytes
        try:
            raw = zf.read(entry_name)
        except Exception as e:
            archive_lineage_rows.append(ArchiveLineageRow(
                archive_lineage_id=alid, job_id=job_id,
                source_filename=os.path.basename(entry_name),
                parent_archive=zip_filename, root_archive=root_archive or zip_filename,
                archive_entry_name=entry_name, extracted_path=safe_path,
                nested_level=current_depth, file_size_bytes=info.file_size,
                extraction_status="EXTRACTION_ERROR", insert_timestamp=ts,
            ))
            failed_files.append(FailedFile(
                source_filename=entry_name, exception_type="ARCHIVE_ERROR",
                reason=f"Could not read ZIP entry: {e}",
                archive_lineage_id=alid,
            ))
            continue

        leaf_name = os.path.basename(entry_name)

        # Nested ZIP
        if entry_name.lower().endswith(".zip"):
            if ing.get("enable_nested_archives", True):
                archive_lineage_rows.append(ArchiveLineageRow(
                    archive_lineage_id=alid, job_id=job_id,
                    source_filename=leaf_name,
                    parent_archive=zip_filename, root_archive=root_archive or zip_filename,
                    archive_entry_name=entry_name, extracted_path=safe_path,
                    nested_level=current_depth, file_size_bytes=info.file_size,
                    extraction_status="SUCCESS", insert_timestamp=ts,
                ))
                _unpack_zip(
                    raw, leaf_name, cfg, job_id,
                    parsed_files, failed_files, archive_lineage_rows,
                    parent_archive=zip_filename,
                    root_archive=root_archive or zip_filename,
                    current_depth=current_depth + 1,
                    file_count_tracker=file_count_tracker,
                    extract_dir=extract_dir,
                )
            else:
                failed_files.append(FailedFile(
                    source_filename=leaf_name, exception_type="ARCHIVE_ERROR",
                    reason="Nested archives are disabled in config",
                ))
            continue

        # Unsupported file type — skip silently but log
        if not _is_supported_extension(entry_name, cfg):
            archive_lineage_rows.append(ArchiveLineageRow(
                archive_lineage_id=alid, job_id=job_id,
                source_filename=leaf_name,
                parent_archive=zip_filename, root_archive=root_archive or zip_filename,
                archive_entry_name=entry_name, extracted_path=safe_path,
                nested_level=current_depth, file_size_bytes=info.file_size,
                extraction_status="BLOCKED_EXTENSION", insert_timestamp=ts,
            ))
            failed_files.append(FailedFile(
                source_filename=leaf_name, exception_type="UNSUPPORTED_FILE_TYPE",
                reason=f"File type not supported in POC: {Path(entry_name).suffix}",
                archive_lineage_id=alid,
            ))
            continue

        # Parse the file
        archive_lineage_rows.append(ArchiveLineageRow(
            archive_lineage_id=alid, job_id=job_id,
            source_filename=leaf_name,
            parent_archive=zip_filename, root_archive=root_archive or zip_filename,
            archive_entry_name=entry_name, extracted_path=safe_path,
            nested_level=current_depth, file_size_bytes=info.file_size,
            extraction_status="SUCCESS", insert_timestamp=ts,
        ))
        try:
            df, fmt_label = _parse_single_file(raw, leaf_name, cfg)
            _, enc = _detect_encoding(raw, cfg) if fmt_label not in ("pdf_text", "pdf_scanned_ocr", "xlsx", "xls", "docx") else (None, None)
            parsed_files.append(ParsedFile(
                source_filename=leaf_name,
                dataframe=df,
                file_format=fmt_label,
                encoding_used=enc,
                archive_lineage_id=alid,
                archive_meta={
                    "parent_archive": zip_filename,
                    "root_archive": root_archive or zip_filename,
                    "nested_level": current_depth,
                    "archive_entry_name": entry_name,
                }
            ))
        except Exception as e:
            exc_type = "PARSE_ERROR"
            if "encoding" in str(e).lower() or "UnicodeDecodeError" in str(type(e).__name__):
                exc_type = "ENCODING_ERROR"
            elif "Unsupported" in str(e):
                exc_type = "UNSUPPORTED_FILE_TYPE"
            failed_files.append(FailedFile(
                source_filename=leaf_name, exception_type=exc_type,
                reason=str(e), archive_lineage_id=alid,
            ))


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------

def parse_input_file(
    file_path: str | None,
    cfg: dict,
    job_id: str,
    raw_bytes: bytes | None = None,
    filename_override: str | None = None,
) -> tuple[list[ParsedFile], list[FailedFile], list[ArchiveLineageRow]]:
    """
    Main entry point. Accepts either a file path or raw bytes.
    Returns (parsed_files, failed_files, archive_lineage_rows).
    """
    parsed_files: list[ParsedFile] = []
    failed_files: list[FailedFile] = []
    archive_lineage_rows: list[ArchiveLineageRow] = []

    if file_path:
        filename = filename_override or os.path.basename(file_path)
        with open(file_path, "rb") as fh:
            raw = fh.read()
    elif raw_bytes is not None:
        raw = raw_bytes
        filename = filename_override or "upload"
    else:
        raise ValueError("Either file_path or raw_bytes must be provided")

    max_total_mb = cfg.get("ingestion", {}).get("max_uncompressed_size_mb", 2048)
    if len(raw) > max_total_mb * 1024 * 1024:
        failed_files.append(FailedFile(
            source_filename=filename, exception_type="ARCHIVE_ERROR",
            reason=f"File size exceeds max_uncompressed_size_mb ({max_total_mb}MB)",
        ))
        return parsed_files, failed_files, archive_lineage_rows

    # ZIP path
    if filename.lower().endswith(".zip"):
        if not cfg.get("ingestion", {}).get("accept_archives", True):
            failed_files.append(FailedFile(
                source_filename=filename, exception_type="UNSUPPORTED_FILE_TYPE",
                reason="Archive ingestion is disabled in config",
            ))
            return parsed_files, failed_files, archive_lineage_rows

        with tempfile.TemporaryDirectory() as tmpdir:
            _unpack_zip(
                raw, filename, cfg, job_id,
                parsed_files, failed_files, archive_lineage_rows,
                parent_archive=None,
                root_archive=filename,
                current_depth=1,
                file_count_tracker=[0],
                extract_dir=tmpdir,
            )
        return parsed_files, failed_files, archive_lineage_rows

    # Blocked extension
    if _is_blocked_extension(filename, cfg):
        failed_files.append(FailedFile(
            source_filename=filename, exception_type="ARCHIVE_ERROR",
            reason=f"Blocked file extension: {Path(filename).suffix}",
        ))
        return parsed_files, failed_files, archive_lineage_rows

    # Unsupported extension
    if not _is_supported_extension(filename, cfg):
        failed_files.append(FailedFile(
            source_filename=filename, exception_type="UNSUPPORTED_FILE_TYPE",
            reason=f"File type not supported in POC: {Path(filename).suffix}",
        ))
        return parsed_files, failed_files, archive_lineage_rows

    # Direct file parse
    try:
        df, fmt_label = _parse_single_file(raw, filename, cfg)
        enc = None
        if fmt_label not in ("pdf_text", "pdf_scanned_ocr", "xlsx", "xls", "docx"):
            _, enc = _detect_encoding(raw, cfg)
        parsed_files.append(ParsedFile(
            source_filename=filename, dataframe=df,
            file_format=fmt_label, encoding_used=enc,
        ))
    except Exception as e:
        exc_type = "PARSE_ERROR"
        if "encoding" in str(e).lower() or "UnicodeDecodeError" in str(type(e).__name__):
            exc_type = "ENCODING_ERROR"
        elif "Unsupported" in str(e):
            exc_type = "UNSUPPORTED_FILE_TYPE"
        failed_files.append(FailedFile(
            source_filename=filename, exception_type=exc_type, reason=str(e),
        ))

    return parsed_files, failed_files, archive_lineage_rows
